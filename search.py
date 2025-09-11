import json
import subprocess
import sys
import time
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
import threading

# Импорты для работы с файлами
import pandas as pd
import docx
import fitz  # PyMuPDF
import requests

# Для GUI
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox

CONFIG_FILE = Path('config.json')
OUTPUT_FILE = Path('data.json')
API_KEY_FILE = Path('api.txt')  # Файл с API-ключом для OpenRouter


class DocumentExtractor:
    """Класс для извлечения текста из различных типов документов"""

    @staticmethod
    def extract_from_word(file_path: Path) -> str:
        """Извлечение текста из Word документов включая таблицы"""
        try:
            doc = docx.Document(file_path)
            text = []

            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Пропускаем пустые параграфы
                    text.append(paragraph.text)

            # Извлекаем текст из таблиц
            # Внимание: python-docx может дублировать текст при объединениях ячеек — здесь базовое извлечение
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ').replace('\t', ' ')
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(' | '.join(row_text))
                if table_text:
                    text.append('\n'.join(table_text))

            return '\n'.join(text)
        except Exception as e:
            return f"Ошибка при извлечении из Word файла {file_path}: {e}"

    @staticmethod
    def extract_from_excel(file_path: Path) -> str:
        """Извлечение текста из Excel файлов"""
        try:
            all_text = []
            with pd.ExcelFile(file_path) as excel_file:
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    sheet_text = df.astype(str).values.flatten()
                    all_text.extend([cell for cell in sheet_text if cell != 'nan'])
            return '\n'.join(all_text)
        except Exception as e:
            return f"Ошибка при извлечении из Excel файла {file_path}: {e}"

    @staticmethod
    def extract_from_pdf(file_path: Path) -> str:
        """Извлечение текста из PDF файлов с помощью PyMuPDF"""
        try:
            texts = []
            with fitz.open(file_path) as doc:
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    # Используем явный режим "text" для более стабильного потока
                    texts.append(page.get_text("text"))
            return '\n'.join(texts)
        except Exception as e:
            return f"Ошибка при извлечении из PDF файла {file_path}: {e}"


class AIInterface:
    """Класс для взаимодействия с AI (Ollama или OpenRouter)"""

    def __init__(self, provider: str = "ollama", api_key: str = None, model: Optional[str] = None):
        self.provider = provider
        self.api_key = api_key
        self.base_url = "http://localhost:11434" if provider == "ollama" else "https://openrouter.ai/api/v1"
        self.model = model if model else ("qwen3:14b" if provider == "ollama" else "deepseek/deepseek-r1:free")
        self.ollama_process = None

        if provider == "openrouter" and not api_key:
            raise ValueError("API-ключ для OpenRouter не предоставлен")

    def clean_model_response(self, response: str) -> str:
        """Очистка ответа модели от лишнего текста (логика сохранена по требованию)"""
        if not response:
            return "null"

        response = response.strip()
        response = re.sub(r'<think>.*?</think>', '', response, flags=re.DOTALL)
        response = response.strip()

        if not response:
            return "null"

        lines = response.split('\n')
        clean_line = ""

        for line in lines:
            line = line.strip()
            if line and not line.startswith('<') and not line.startswith('Объяснение'):
                clean_line = line
                break

        if not clean_line:
            return "null"

        clean_line = clean_line.strip('"').strip("'")

        explanation_patterns = [
            'найдено', 'содержится', 'указано', 'упоминается',
            'в тексте', 'анализ', 'рассмотрим', 'видно что',
            'согласно', 'поэтому', 'таким образом'
        ]

        if any(pattern in clean_line.lower() for pattern in explanation_patterns):
            return "null"

        if len(clean_line) > 500:
            return "null"

        return clean_line

    def start(self, logger=None) -> bool:
        """Запуск провайдера (только для Ollama)"""
        if self.provider != "ollama":
            return True

        def log(msg):
            if logger:
                logger(msg)

        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            if response.status_code == 200:
                log("Ollama уже запущена")
                return True
        except requests.exceptions.RequestException as e:
            log(f"Ollama не обнаружена по HTTP: {e}")

        try:
            self.ollama_process = subprocess.Popen(
                ["ollama", "serve"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )

            max_attempts = 30
            for attempt in range(1, max_attempts + 1):
                try:
                    response = requests.get(f"{self.base_url}/api/tags", timeout=2)
                    if response.status_code == 200:
                        log("Ollama успешно запущена")
                        return True
                    else:
                        log(f"Ожидание Ollama... код {response.status_code}")
                except requests.exceptions.RequestException:
                    log(f"Ожидание Ollama... попытка {attempt}/{max_attempts}")
                time.sleep(2)

            log("Не удалось дождаться запуска Ollama в отведенное время")
            return False

        except Exception as e:
            if logger:
                logger(f"Ошибка запуска Ollama: {e}")
            return False

    def query_model(self, text: str, keywords: List[str], logger=None) -> str:
        """Запрос к модели для поиска значений"""
        keywords_str = ", ".join(keywords)
        prompt = f"""Представь что ты робот-парсер твоя задача найти "{keywords_str}" в тексте: "{text}". Строжайше выводи только то значение которое у тебя запрашивают так как твои значения используются в программе и лишний текст будет ей мешать. """

        def log(msg):
            if logger:
                logger(msg)

        if self.provider == "ollama":
            try:
                payload = {
                    "model": self.model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,
                        "top_p": 0.9,
                        "think": False
                    }
                }

                response = requests.post(
                    f"{self.base_url}/api/generate",
                    json=payload,
                    timeout=300
                )

                if response.status_code == 200:
                    result = response.json()
                    raw_answer = result.get("response", "").strip()
                    cleaned_answer = self.clean_model_response(raw_answer)
                    return cleaned_answer if cleaned_answer else "null"
                else:
                    log(f"ОШИБКА Ollama: status={response.status_code}, body={response.text[:300]}")
                    return "null"

            except Exception as e:
                log(f"Исключение при запросе к Ollama: {e}")
                return "null"

        elif self.provider == "openrouter":
            try:
                headers = {
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "http://localhost",
                    "X-Title": "Document AI Parser"
                }
                payload = {
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "max_tokens": 500
                }

                response = requests.post(
                    f"{self.base_url}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=300
                )

                if response.status_code == 200:
                    result = response.json()
                    raw_answer = result["choices"][0]["message"]["content"].strip()
                    cleaned_answer = self.clean_model_response(raw_answer)
                    return cleaned_answer if cleaned_answer else "null"
                else:
                    log(f"ОШИБКА OpenRouter: status={response.status_code}, body={response.text[:300]}")
                    return "null"
            except Exception as e:
                log(f"Исключение при запросе к OpenRouter: {e}")
                return "null"

    def stop(self):
        """Остановка провайдера (только для Ollama)"""
        if self.provider == "ollama" and self.ollama_process:
            try:
                self.ollama_process.terminate()
                self.ollama_process.wait(timeout=10)
            except Exception:
                pass
            finally:
                self.ollama_process = None


class DocumentProcessor:
    """Основной класс для обработки документов"""

    def __init__(self, gui_log):
        self.extractor = DocumentExtractor()
        self.ai = None
        self.not_found_items = []
        self.gui_log = gui_log

    def set_ai_interface(self, ai_interface: AIInterface):
        self.ai = ai_interface

    def load_config(self) -> Optional[Dict]:
        """Загрузка конфигурации"""
        if not CONFIG_FILE.exists():
            self.gui_log(f"Файл конфигурации {CONFIG_FILE} не найден")
            return None

        try:
            with CONFIG_FILE.open('r', encoding='utf-8') as f:
                config = json.load(f)
            return config
        except Exception as e:
            self.gui_log(f"Ошибка при загрузке конфигурации: {e}")
            return None

    def _safe_join_under_root(self, root: Path, rel: str) -> Optional[Path]:
        """Безопасно соединяет root и относительный путь, не позволяя выйти за root"""
        try:
            candidate = (root / rel).resolve()
            root_resolved = root.resolve()
            if hasattr(candidate, "is_relative_to"):
                # Python 3.9+
                if candidate.is_relative_to(root_resolved):
                    return candidate
                else:
                    return None
            else:
                # Совместимость для Python < 3.9
                candidate_str = str(candidate)
                root_str = str(root_resolved)
                if candidate_str.startswith(root_str):
                    return candidate
                return None
        except Exception:
            return None

    def extract_text_from_file(self, file_path: Path, file_type: str) -> str:
        """Извлечение текста из файла в зависимости от типа"""
        if not file_path.exists():
            self.gui_log(f"Файл не найден: {file_path}")
            return ""

        if file_type == "word":
            return self.extractor.extract_from_word(file_path)
        elif file_type == "excel":
            return self.extractor.extract_from_excel(file_path)
        elif file_type == "pdf":
            return self.extractor.extract_from_pdf(file_path)
        else:
            self.gui_log(f"Неподдерживаемый тип файла: {file_type}")
            return ""

    def process_documents(self) -> List[Dict[str, Any]]:
        """Основная функция обработки всех документов"""
        # Сбрасываем список не найденных элементов перед запуском
        self.not_found_items = []

        config = self.load_config()
        if not config:
            return []

        root_value = config.get('root', '')
        if not root_value:
            self.gui_log("В конфигурации не указан корневой путь 'root'")
            return []

        root_path = Path(root_value)
        if not root_path.exists() or not root_path.is_dir():
            self.gui_log(f"Корневая папка не найдена или не является папкой: {root_path}")
            return []

        # Проверяем, что AI интерфейс установлен
        if not self.ai:
            self.gui_log("AI интерфейс не установлен")
            return []

        results = []
        items = config.get('items', [])

        self.gui_log(f"Обработка {len(items)} элементов...")

        for i, item in enumerate(items, 1):
            data_name = item.get('data_name', '')
            relative_file_path = item.get('file', '')
            file_type = item.get('type', '')
            keywords = item.get('keywords', [])

            self.gui_log(f"\n[{i}/{len(items)}] Обработка: {data_name}")

            full_file_path = None
            if relative_file_path:
                full_file_path = self._safe_join_under_root(root_path, relative_file_path)
                if full_file_path is None:
                    self.gui_log(f"  ❌ Недопустимый путь (вылазка за root): {relative_file_path}")
            else:
                self.gui_log(f"  ❌ Файл не указан")

            status = 'not_found'
            ai_result = "null"
            reason = ''

            if not full_file_path or not (full_file_path.exists() and full_file_path.is_file()):
                reason = 'Файл не указан или не найден'
                self.not_found_items.append({
                    'data_name': data_name,
                    'file': relative_file_path,
                    'reason': reason,
                    'keywords': keywords
                })
            else:
                text = self.extract_text_from_file(full_file_path, file_type)
                if not text.strip() or text.startswith("Ошибка при извлечении"):
                    self.gui_log(f"  ❌ Не удалось извлечь текст из: {relative_file_path}")
                    reason = 'Не удалось извлечь текст'
                    self.not_found_items.append({
                        'data_name': data_name,
                        'file': relative_file_path,
                        'reason': reason,
                        'keywords': keywords
                    })
                else:
                    if keywords:
                        self.gui_log(f"  🔍 Поиск ключевых слов: {keywords}")
                        ai_result = self.ai.query_model(text, keywords, logger=self.gui_log)
                        if ai_result == "null" or not ai_result:
                            self.gui_log(f"  ❌ Значение не найдено")
                            reason = 'Нейросеть не нашла значение'
                            self.not_found_items.append({
                                'data_name': data_name,
                                'file': relative_file_path,
                                'keywords': keywords,
                                'reason': reason
                            })
                        else:
                            self.gui_log(f"  ✅ Найдено: {ai_result[:100]}...")
                            status = 'found'
                    else:
                        self.gui_log(f"  ❌ Ключевые слова не указаны")
                        reason = 'Ключевые слова не указаны'
                        self.not_found_items.append({
                            'data_name': data_name,
                            'file': relative_file_path,
                            'reason': reason
                        })

            # Сохраняем результат по элементу (одна запись на элемент)
            result = {
                'data_name': data_name,
                'file': relative_file_path,
                'type': file_type,
                'keywords': keywords,
                'extracted_value': ai_result,
                'status': status
            }
            if status != 'found' and reason:
                result['reason'] = reason
            results.append(result)

        # ВНИМАНИЕ: больше не добавляем self.not_found_items в results повторно — дубликатов не будет
        return results

    def save_results(self, results: List[Dict[str, Any]]):
        """Сохранение результатов в JSON файл"""
        try:
            with OUTPUT_FILE.open('w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=4)
            self.gui_log(f"\n✅ Результаты сохранены в {OUTPUT_FILE}")
        except Exception as e:
            self.gui_log(f"❌ Ошибка при сохранении результатов: {e}")

    def print_report(self, results: List[Dict[str, Any]]):
        """Вывод отчета о результатах"""
        total = len(results)
        found = sum(1 for r in results if r['status'] == 'found')
        not_found = total - found

        self.gui_log("\n" + "=" * 50)
        self.gui_log("📊 ОТЧЕТ О РЕЗУЛЬТАТАХ")
        self.gui_log("=" * 50)
        self.gui_log(f"Всего обработано: {total}")
        self.gui_log(f"Найдено значений: {found}")
        self.gui_log(f"Не найдено значений: {not_found}")

        if self.not_found_items:
            self.gui_log(f"\n❌ НЕ НАЙДЕННЫЕ ЗНАЧЕНИЯ ({len(self.not_found_items)}):")
            self.gui_log("-" * 50)
            for i, item in enumerate(self.not_found_items, 1):
                self.gui_log(f"{i}. {item.get('data_name', '')}")
                self.gui_log(f"   Файл: {item.get('file', 'не указан')}")
                self.gui_log(f"   Причина: {item.get('reason', '')}")
                if 'keywords' in item and item['keywords']:
                    self.gui_log(f"   Ключевые слова: {item['keywords']}")
                self.gui_log("")


class GUIApp:
    """Графический интерфейс на Tkinter"""

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Document AI Parser")
        self.root.geometry("700x550")

        self.processor = DocumentProcessor(self.print_to_log)

        # Выбор провайдера
        ttk.Label(self.root, text="Выберите AI-провайдера:").pack(pady=10)
        self.provider_var = tk.StringVar(value="ollama")

        # Добавляем обработчики изменения провайдера
        ollama_radio = ttk.Radiobutton(self.root, text="Локальная Ollama",
                                       variable=self.provider_var, value="ollama",
                                       command=self.update_default_model)
        ollama_radio.pack()

        openrouter_radio = ttk.Radiobutton(self.root, text="OpenRouter API (ключ из api.txt)",
                                           variable=self.provider_var, value="openrouter",
                                           command=self.update_default_model)
        openrouter_radio.pack()

        # Ввод модели
        ttk.Label(self.root, text="Название модели (опционально):").pack(pady=5)
        self.model_var = tk.StringVar()
        ttk.Entry(self.root, textvariable=self.model_var).pack(fill='x', padx=10)

        # Устанавливаем модель по умолчанию для начального провайдера
        self.update_default_model()

        # Кнопка запуска
        self.start_button = ttk.Button(self.root, text="Запустить обработку", command=self.start_processing)
        self.start_button.pack(pady=20)

        # Область логов
        self.log_text = scrolledtext.ScrolledText(self.root, height=20, width=85)
        self.log_text.pack(pady=10, padx=10, fill='both', expand=True)

        self.root.mainloop()

    def update_default_model(self):
        """Обновляет модель по умолчанию в зависимости от выбранного провайдера"""
        provider = self.provider_var.get()
        if provider == "ollama":
            default_model = "qwen3:14b"
        else:  # openrouter
            default_model = "deepseek/deepseek-r1:free"

        # Устанавливаем модель по умолчанию только если поле пустое
        if not self.model_var.get():
            self.model_var.set(default_model)

    def print_to_log(self, text: str):
        """Вывод в лог GUI"""
        try:
            self.log_text.insert(tk.END, text + "\n")
            self.log_text.see(tk.END)
            self.root.update_idletasks()
        except Exception:
            # На случай, если GUI уже закрыт
            pass

    def start_processing(self):
        """Запуск обработки в отдельном потоке"""
        self.start_button.config(state='disabled')
        threading.Thread(target=self.run_processing, daemon=True).start()

    def run_processing(self):
        """Логика обработки"""
        provider = self.provider_var.get()
        api_key = None
        model = self.model_var.get() or None
        ai = None

        if provider == "openrouter":
            # Чтение API-ключа из api.txt
            if not API_KEY_FILE.exists():
                msg = f"Файл {API_KEY_FILE} не найден. Создайте его с API-ключом."
                self.print_to_log(f"Ошибка: {msg}")
                try:
                    messagebox.showerror("Ошибка", msg)
                except Exception:
                    pass
                self.start_button.config(state='normal')
                return

            try:
                with API_KEY_FILE.open('r', encoding='utf-8') as f:
                    api_key = f.read().strip()
                if not api_key:
                    msg = f"Файл {API_KEY_FILE} пустой. Добавьте API-ключ."
                    self.print_to_log(f"Ошибка: {msg}")
                    try:
                        messagebox.showerror("Ошибка", msg)
                    except Exception:
                        pass
                    self.start_button.config(state='normal')
                    return
            except Exception as e:
                msg = f"Не удалось прочитать {API_KEY_FILE}: {str(e)}"
                self.print_to_log(msg)
                try:
                    messagebox.showerror("Ошибка", msg)
                except Exception:
                    pass
                self.start_button.config(state='normal')
                return

        try:
            ai = AIInterface(provider=provider, api_key=api_key, model=model)
            if not ai.start(logger=self.print_to_log):
                self.print_to_log("Не удалось запустить AI-провайдера")
                try:
                    messagebox.showerror("Ошибка", "Не удалось запустить AI-провайдера")
                except Exception:
                    pass
                self.start_button.config(state='normal')
                return

            self.processor.set_ai_interface(ai)
            self.print_to_log("🚀 Запуск обработки...")

            results = self.processor.process_documents()
            self.processor.save_results(results)
            self.processor.print_report(results)

            self.print_to_log("✅ Обработка завершена")
        except Exception as e:
            self.print_to_log(f"❌ Ошибка: {str(e)}")
            try:
                messagebox.showerror("Ошибка", str(e))
            except Exception:
                pass
        finally:
            if ai:
                ai.stop()
            self.start_button.config(state='normal')


if __name__ == "__main__":
    GUIApp()
